from docx import Document

def generate_docx_report(title, summary_text, charts, output_path="report.docx"):
    doc = Document()
    doc.add_heading(title, 0)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary_text)

    for i, chart_path in enumerate(charts):
        doc.add_heading(f"Chart {i+1}", level=2)
        doc.add_picture(chart_path, width=docx.shared.Inches(5))

    doc.save(output_path)
    return output_path
